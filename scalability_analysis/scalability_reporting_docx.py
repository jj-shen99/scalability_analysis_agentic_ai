#!/usr/bin/env python3
"""
Scalability Reporting Module - DOCX Generator

This module provides functionality for generating DOCX reports from
scalability analysis results with enhanced formatting and styling.

Part of the modular scalability analysis framework
"""

import os
from datetime import datetime
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor


def generate_docx_report(analysis_results, output_dir, plot_paths=None):
    """
    Generate a DOCX report from scalability analysis results
    
    Args:
        analysis_results (list): List of dictionaries with analysis results
        output_dir (str): Directory to save the report
        plot_paths (dict, optional): Dictionary of paths to plot images
        
    Returns:
        str: Path to the generated report
    """
    # Sort results by resource level for consistent reporting
    sorted_results = sorted(analysis_results, key=lambda x: x.get('resource_level', 0))
    
    # Extract resource levels and metrics for comparison
    resource_levels = [r.get('resource_level', 0) for r in sorted_results]
    throughputs = [r.get('metrics', {}).get('throughput', 0) for r in sorted_results]
    response_times = [r.get('metrics', {}).get('avg_response_time', 0) for r in sorted_results]
    
    # Calculate speedups relative to the baseline
    baseline_throughput = throughputs[0] if throughputs else 1
    speedups = [t / baseline_throughput for t in throughputs]

    # Get scalability analysis if available
    scalability_analysis = {}
    for result in sorted_results:
        if 'scalability_analysis' in result:
            scalability_analysis = result['scalability_analysis']
            break

    # Create a new Document
    doc = docx.Document()
    
    # Add document properties
    doc.core_properties.title = "Scalability Analysis Report"
    doc.core_properties.subject = "Performance Scalability"
    doc.core_properties.created = datetime.now()

    # Document styles setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add title
    title = doc.add_heading('Scalability Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_run.italic = True
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Executive Summary section
    doc.add_heading('Executive Summary', 1)
    
    if len(resource_levels) > 1:
        summary = doc.add_paragraph(
            f"This report analyzes the scalability characteristics of the system under test across {len(resource_levels)} "
            f"different resource levels (from {min(resource_levels)} to {max(resource_levels)})."
        )
    else:
        summary = doc.add_paragraph(
            "This report analyzes the performance characteristics of the system under test."
        )
    
    # Add Key Findings section
    doc.add_heading('Key Findings', 2)
    
    findings_list = []
    
    if throughputs:
        max_throughput = max(throughputs)
        max_throughput_level = resource_levels[throughputs.index(max_throughput)]
        findings_list.append(f"Maximum Throughput: {max_throughput:.2f} requests/sec achieved at resource level {max_throughput_level}")
    
    if response_times:
        min_response_time = min(response_times)
        min_response_time_level = resource_levels[response_times.index(min_response_time)]
        findings_list.append(f"Best Response Time: {min_response_time:.2f} ms achieved at resource level {min_response_time_level}")
    
    if len(speedups) > 1:
        max_speedup = max(speedups[1:], default=1)  # Skip baseline
        max_speedup_level = resource_levels[speedups.index(max_speedup)]
        findings_list.append(f"Maximum Speedup: {max_speedup:.2f}x achieved at resource level {max_speedup_level} compared to baseline")
    
    # Add findings as bullet points
    for finding in findings_list:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(finding)
    
    # Add Detailed Performance Metrics section
    doc.add_heading('Detailed Performance Metrics', 1)
    
    # Add table for metrics
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Resource Level'
    header_cells[1].text = 'Throughput (req/s)'
    header_cells[2].text = 'Avg Response Time (ms)'
    header_cells[3].text = 'Error %'
    
    # Apply header formatting
    for cell in header_cells:
        cell_para = cell.paragraphs[0]
        cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell_para.runs:
            run.bold = True
    
    # Add data rows
    for result in sorted_results:
        metrics = result.get('metrics', {})
        level = result.get('resource_level', 0)
        throughput = metrics.get('throughput', 0)
        response_time = metrics.get('avg_response_time', 0)
        error_percentage = metrics.get('error_percentage', 0)
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(level)
        row_cells[1].text = f"{throughput:.2f}"
        row_cells[2].text = f"{response_time:.2f}"
        row_cells[3].text = f"{error_percentage:.2f}"
    
    # Add Basic Scalability Metrics section if we have multiple resource levels
    if len(resource_levels) > 1:
        doc.add_heading('Basic Scalability Metrics', 1)
        
        # Add table for scalability metrics
        scalability_table = doc.add_table(rows=1, cols=3)
        scalability_table.style = 'Table Grid'
        
        # Add header row
        header_cells = scalability_table.rows[0].cells
        header_cells[0].text = 'Resource Level'
        header_cells[1].text = 'Throughput Speedup'
        header_cells[2].text = 'Scalability Efficiency'
        
        # Apply header formatting
        for cell in header_cells:
            cell_para = cell.paragraphs[0]
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell_para.runs:
                run.bold = True
        
        # Add data rows
        for i, level in enumerate(resource_levels):
            speedup = speedups[i]
            # Calculate efficiency as speedup / resource_level relative to baseline
            relative_resources = level / resource_levels[0]
            efficiency = speedup / relative_resources if relative_resources > 0 else 0
            
            row_cells = scalability_table.add_row().cells
            row_cells[0].text = str(level)
            row_cells[1].text = f"{speedup:.2f}x"
            row_cells[2].text = f"{efficiency:.2%}"
    
    # Add Advanced Scalability Analysis section if available
    if scalability_analysis and 'models' in scalability_analysis:
        doc.add_heading('Advanced Scalability Analysis', 1)
        models = scalability_analysis.get('models', {})
        
        # Amdahl's Law section
        if 'amdahl' in models:
            p = models['amdahl']
            max_speedup = 1 / (1 - p) if p < 1 else float('inf')
            
            doc.add_heading("Amdahl's Law Analysis", 2)
            
            amdahl_para = doc.add_paragraph()
            amdahl_para.add_run(f"Parallelizable portion: ").bold = True
            amdahl_para.add_run(f"{p:.2%}")
            
            amdahl_para = doc.add_paragraph()
            amdahl_para.add_run(f"Serial portion: ").bold = True
            amdahl_para.add_run(f"{1-p:.2%}")
            
            amdahl_para = doc.add_paragraph()
            amdahl_para.add_run(f"Theoretical maximum speedup: ").bold = True
            amdahl_para.add_run(f"{max_speedup:.2f}x")
        
        # Gustafson's Law section
        if 'gustafson' in models:
            p = models['gustafson']
            
            doc.add_heading("Gustafson's Law Analysis", 2)
            
            gustafson_para = doc.add_paragraph()
            gustafson_para.add_run(f"Scalable portion: ").bold = True
            gustafson_para.add_run(f"{p:.2%}")
            
            gustafson_para = doc.add_paragraph()
            gustafson_para.add_run(f"Fixed portion: ").bold = True
            gustafson_para.add_run(f"{1-p:.2%}")
        
        # USL section
        if 'usl' in models:
            sigma, kappa = models['usl']
            
            doc.add_heading("Universal Scalability Law Analysis", 2)
            
            usl_para = doc.add_paragraph()
            usl_para.add_run(f"Contention factor (σ): ").bold = True
            usl_para.add_run(f"{sigma:.4f}")
            
            usl_para = doc.add_paragraph()
            usl_para.add_run(f"Coherency factor (κ): ").bold = True
            usl_para.add_run(f"{kappa:.4f}")
            
            # Calculate peak concurrency point
            if kappa > 0:
                peak_n = (1 - sigma) / (2 * kappa) if (1 - sigma) > 0 else 1
                
                usl_para = doc.add_paragraph()
                usl_para.add_run(f"Optimal concurrency: ").bold = True
                usl_para.add_run(f"{peak_n:.2f} resources")
        
        # Add interpretations if available
        if 'interpretations' in scalability_analysis:
            interpretations = scalability_analysis.get('interpretations', {})
            
            doc.add_heading("Model Interpretations", 2)
            
            for model_name, interp in interpretations.items():
                if 'assessment' in interp:
                    interp_para = doc.add_paragraph(style='List Bullet')
                    interp_para.add_run(f"{model_name.capitalize()}: ").bold = True
                    interp_para.add_run(interp['assessment'])
    
    # Add Visual Analysis section
    doc.add_heading('Visual Analysis', 1)
    doc.add_paragraph("The following plots provide visual representation of the scalability characteristics.")
    
    # Add plots if available
    if plot_paths:
        # Throughput plot
        if 'throughput' in plot_paths:
            doc.add_heading("Throughput vs. Resource Level", 2)
            doc.add_picture(plot_paths['throughput'], width=Inches(6.0))
            doc.add_paragraph("This plot shows how the system throughput changes as resources are added. "
                             "The trend line indicates the overall scaling pattern.")
        
        # Response time plot
        if 'response_time' in plot_paths:
            doc.add_heading("Response Time vs. Resource Level", 2)
            doc.add_picture(plot_paths['response_time'], width=Inches(6.0))
            doc.add_paragraph("This plot illustrates how response times are affected by resource scaling. "
                             "Lower values indicate better performance.")
        
        # Speedup plot
        if 'speedup' in plot_paths:
            doc.add_heading("Speedup vs. Resource Level", 2)
            doc.add_picture(plot_paths['speedup'], width=Inches(6.0))
            doc.add_paragraph("This plot compares actual speedup against ideal linear speedup. "
                             "The gap between actual and ideal lines indicates efficiency loss as resources scale.")
        
        # Model comparison plot if available
        if 'models_comparison' in plot_paths:
            doc.add_heading("Scalability Models Comparison", 2)
            doc.add_picture(plot_paths['models_comparison'], width=Inches(6.0))
            doc.add_paragraph("This plot compares the actual speedup with predictions from different scalability laws. "
                             "The closest model to actual data points indicates which theoretical model best describes the system's scaling behavior.")
        
        # Theoretical projections plot if available
        if 'theoretical_projections' in plot_paths:
            doc.add_heading("Theoretical Scalability Projections", 2)
            doc.add_picture(plot_paths['theoretical_projections'], width=Inches(6.0))
            doc.add_paragraph("This plot shows theoretical projections of different scalability models based on observed data. "
                             "It predicts how the system might scale with additional resources beyond those tested.")
        
        # Add model characteristics plot if available
        if 'model_characteristics' in plot_paths:
            doc.add_heading('Comparative Scalability Model Characteristics', 2)
            p = doc.add_paragraph('This educational plot illustrates the fundamental differences between various scalability models with different parameters. ')
            p.add_run('It helps identify which theoretical model best describes your system\'s behavior based on the shape of your actual data points.')
            
            p2 = doc.add_paragraph()
            p2.add_run('Key model characteristics:').bold = True
            models_list = doc.add_paragraph(style='List Bullet')
            models_list.add_run('Linear scaling: ').bold = True
            models_list.add_run('Represents perfect scaling where adding resources results in proportional performance improvement.')
            
            models_list2 = doc.add_paragraph(style='List Bullet')
            models_list2.add_run('Amdahl\'s Law: ').bold = True
            models_list2.add_run('Shows diminishing returns due to serial portions in the workload, with a clear asymptotic limit.')
            
            models_list3 = doc.add_paragraph(style='List Bullet')
            models_list3.add_run('Gustafson\'s Law: ').bold = True
            models_list3.add_run('Exhibits better scaling than Amdahl\'s when workload size increases with resources.')
            
            models_list4 = doc.add_paragraph(style='List Bullet')
            models_list4.add_run('Universal Scalability Law: ').bold = True
            models_list4.add_run('Shows initial scaling followed by performance degradation due to contention and coherency delays.')
            
            doc.add_picture(plot_paths['model_characteristics'], width=Inches(6.0))
            doc.add_paragraph("This educational plot illustrates the fundamental differences between various scalability models "
                             "with different parameters. It helps identify which theoretical model best describes your system's behavior.")
        
        # Efficiency analysis plot if available
        if 'efficiency' in plot_paths:
            doc.add_heading("Scalability Efficiency Analysis", 2)
            doc.add_picture(plot_paths['efficiency'], width=Inches(6.0))
            efficiency_para = doc.add_paragraph("This plot shows how efficiently your system scales as resources increase. ")
            efficiency_para.add_run("The efficiency is calculated as speedup divided by resource ratio, with 100% representing perfect linear scaling. ")
            efficiency_para.add_run("Declining efficiency indicates diminishing returns from additional resources.")
            
            # Add key observations about efficiency if there are at least two data points
            if len(resource_levels) >= 2:
                # Calculate baseline efficiency
                baseline_throughput = throughputs[0]
                baseline_resource = resource_levels[0]
                speedups = [t / baseline_throughput for t in throughputs]
                resource_ratios = [r / baseline_resource for r in resource_levels]
                efficiencies = [s / r for s, r in zip(speedups, resource_ratios)]
                
                last_efficiency = efficiencies[-1] if efficiencies else 0
                
                efficiency_obs = doc.add_paragraph(style='List Bullet')
                efficiency_obs.add_run(f"Current scaling efficiency: ").bold = True
                efficiency_obs.add_run(f"{last_efficiency:.1%} at {resource_levels[-1]} resources")
                
                if last_efficiency > 0.9:
                    conclusion = doc.add_paragraph(style='List Bullet')
                    conclusion.add_run("Observation: ").bold = True
                    conclusion.add_run("Excellent scaling efficiency indicates your system has minimal serialization bottlenecks")
                elif last_efficiency > 0.7:
                    conclusion = doc.add_paragraph(style='List Bullet')
                    conclusion.add_run("Observation: ").bold = True
                    conclusion.add_run("Good scaling efficiency indicates your system parallelizes well with some overhead")
                else:
                    conclusion = doc.add_paragraph(style='List Bullet')
                    conclusion.add_run("Observation: ").bold = True
                    conclusion.add_run("Lower scaling efficiency suggests significant serialization or contention in your system")
        
        # Efficiency heatmap if available
        if 'heatmap' in plot_paths:
            doc.add_heading("Efficiency Heatmap", 2)
            doc.add_picture(plot_paths['heatmap'], width=Inches(6.0))
            heatmap_para = doc.add_paragraph("This heatmap visualizes scaling efficiency between different resource levels. ")
            heatmap_para.add_run("Each cell shows the efficiency when scaling from the baseline (row) to the target (column) resource level. ")
            heatmap_para.add_run("Values close to 1.0 (green) indicate good scaling efficiency.")
        
        # Cost efficiency plot if available
        if 'cost_efficiency' in plot_paths:
            doc.add_heading("Cost Efficiency Analysis", 2)
            doc.add_picture(plot_paths['cost_efficiency'], width=Inches(6.0))
            cost_para = doc.add_paragraph("This dual visualization shows the relationship between cost and performance (left) ")
            cost_para.add_run("and which configuration provides the best throughput per cost unit (right). The highlighted bar indicates ")
            cost_para.add_run("the most cost-effective resource level for optimal return on investment.")
            
            # Add more detailed cost analysis
            doc.add_paragraph()
            cost_analysis = doc.add_paragraph()
            cost_analysis.add_run("Why cost efficiency matters: ").bold = True
            cost_analysis.add_run("As you scale your system, it's crucial to understand not just raw performance gains but the cost-to-performance ratio. ")
            cost_analysis.add_run("The linear cost model used in this analysis assumes costs scale directly with resource count.")
            
            # Add resource-specific insights
            if len(resource_levels) >= 2:
                # Find most cost-efficient level
                cost_efficiency = []
                for i, (level, throughput) in enumerate(zip(resource_levels, throughputs)):
                    cost = level  # Linear cost model
                    efficiency = throughput / cost if cost > 0 else 0
                    cost_efficiency.append((level, efficiency))
                
                # Sort by efficiency
                cost_efficiency.sort(key=lambda x: x[1], reverse=True)
                best_level, best_efficiency = cost_efficiency[0]
                
                cost_insight = doc.add_paragraph(style='List Bullet')
                cost_insight.add_run(f"Most cost-efficient configuration: ").bold = True
                cost_insight.add_run(f"{best_level} resources with {best_efficiency:.2f} throughput per cost unit")
                
                # Compare to highest resource level
                if best_level != max(resource_levels):
                    highest_level = max(resource_levels)
                    highest_idx = resource_levels.index(highest_level)
                    highest_efficiency = throughputs[highest_idx] / highest_level
                    
                    cost_compare = doc.add_paragraph(style='List Bullet')
                    cost_compare.add_run("Cost efficiency comparison: ").bold = True
                    cost_compare.add_run(f"The most cost-efficient configuration is {best_efficiency/highest_efficiency:.1%} ")
                    cost_compare.add_run(f"more economical than using {highest_level} resources")
                
                recommendation = doc.add_paragraph(style='List Bullet')
                recommendation.add_run("Recommendation: ").bold = True
                recommendation.add_run(f"For optimal resource utilization, consider using {best_level} resources when budget constraints are a priority.")
            
        # Algorithm complexity plot if available
        if 'algorithm_complexity' in plot_paths:
            doc.add_heading("Algorithm Complexity Analysis", 2)
            doc.add_picture(plot_paths['algorithm_complexity'], width=Inches(6.0))
            
            # Add detailed explanation about algorithm complexity analysis
            algo_para = doc.add_paragraph("This plot shows the algorithmic complexity of your system by fitting various complexity models ")
            algo_para.add_run("(O(1), O(log n), O(n), O(n log n), O(n²), etc.) to your performance data. The best fitting model indicates ")
            algo_para.add_run("your system's computational complexity class, which influences how it will scale with increasing workloads.")
            
            # Add explanation of complexity classes
            doc.add_paragraph()
            algo_complexity = doc.add_paragraph()
            algo_complexity.add_run("Understanding complexity classes: ").bold = True
            
            complexity_o1 = doc.add_paragraph(style='List Bullet')
            complexity_o1.add_run("O(1) - Constant time: ").bold = True
            complexity_o1.add_run("Performance remains constant regardless of input size. Ideal for lookups and cached operations.")
            
            complexity_ologn = doc.add_paragraph(style='List Bullet')
            complexity_ologn.add_run("O(log n) - Logarithmic time: ").bold = True
            complexity_ologn.add_run("Performance increases logarithmically with input size. Common in tree-based structures and binary searches.")
            
            complexity_on = doc.add_paragraph(style='List Bullet')
            complexity_on.add_run("O(n) - Linear time: ").bold = True
            complexity_on.add_run("Performance scales linearly with input size. Typical in operations that process each input element once.")
            
            complexity_onlogn = doc.add_paragraph(style='List Bullet')
            complexity_onlogn.add_run("O(n log n) - Linearithmic time: ").bold = True
            complexity_onlogn.add_run("Performance scales slightly worse than linear. Common in efficient sorting algorithms.")
            
            complexity_on2 = doc.add_paragraph(style='List Bullet')
            complexity_on2.add_run("O(n²) - Quadratic time: ").bold = True
            complexity_on2.add_run("Performance scales with the square of input size. Often indicates nested loops or inefficient algorithms.")
            
            doc.add_paragraph()
            algo_importance = doc.add_paragraph()
            algo_importance.add_run("Why this matters for scalability: ").bold = True
            algo_importance.add_run("The algorithmic complexity of your system directly impacts how it will handle larger workloads. ")
            algo_importance.add_run("Systems with better complexity classes (O(1), O(log n)) will generally scale better than those with higher complexity classes (O(n), O(n²)).")
            
            # Add algorithm complexity interpretation if available
            first_result = sorted_results[0] if sorted_results else None
            if first_result and 'advanced_analysis' in first_result and 'algorithm_complexity' in first_result['advanced_analysis']:
                complexity = first_result['advanced_analysis']['algorithm_complexity']
                if 'interpretation' in complexity and complexity['interpretation'].get('success', False):
                    interpretation = complexity['interpretation']
                    best_model = interpretation.get('best_model', 'Unknown')
                    confidence = interpretation.get('confidence', '')
                    
                    doc.add_heading("Algorithm Complexity Interpretation", 3)
                    
                    model_para = doc.add_paragraph()
                    model_para.add_run(f"Best fitting model: ").bold = True
                    model_para.add_run(f"{best_model}")
                    
                    conf_para = doc.add_paragraph()
                    conf_para.add_run(f"Confidence: ").bold = True
                    conf_para.add_run(f"{confidence}")
                    
                    explain_para = doc.add_paragraph()
                    explain_para.add_run(f"Explanation: ").bold = True
                    explain_para.add_run(f"{interpretation.get('explanation', '')}")
                    
                    impl_para = doc.add_paragraph()
                    impl_para.add_run(f"Implications: ").bold = True
                    impl_para.add_run(f"{interpretation.get('implications', '')}")
                    # Check if load scalability analysis data exists in the first result
                    # Initialize key_metrics with empty dictionary
                    key_metrics = {}
                    
                    if analysis_results and isinstance(analysis_results, list) and len(analysis_results) > 0:
                        # Try to access load scalability analysis data
                        if 'advanced_analysis' in analysis_results[0] and 'load_scalability' in analysis_results[0]['advanced_analysis']:
                            load_data = analysis_results[0]['advanced_analysis']['load_scalability']
                            if load_data and 'interpretation' in load_data:
                                load_interpretation = load_data['interpretation']
                                key_metrics = load_interpretation.get('key_metrics', {})
                                
                                doc.add_heading("Load Scalability Interpretation", 3)
                                
                                sat_para = doc.add_paragraph()
                                sat_para.add_run(f"Saturation point: ").bold = True
                                sat_para.add_run(f"{key_metrics.get('saturation_load', 'Unknown')} users/requests")
                    
                    opt_para = doc.add_paragraph()
                    opt_para.add_run(f"Optimal load point: ").bold = True
                    opt_para.add_run(f"{key_metrics.get('optimal_load', 'Unknown')} users/requests")
                    
                    deg_para = doc.add_paragraph()
                    deg_para.add_run(f"Performance degradation observed: ").bold = True
                    deg_para.add_run(f"{'Yes' if key_metrics.get('has_degradation', False) else 'No'}")
                    
                    if interpretation.get('insights'):
                        doc.add_heading("Key Insights", 4)
                        for insight in interpretation.get('insights', []):
                            insight_para = doc.add_paragraph(style='List Bullet')
                            insight_para.add_run(insight)
                    
                    if interpretation.get('recommendations'):
                        doc.add_heading("Recommendations", 4)
                        for recommendation in interpretation.get('recommendations', []):
                            rec_para = doc.add_paragraph(style='List Bullet')
                            rec_para.add_run(recommendation)
        
        # Load Scalability Analysis if available
        if plot_paths and 'load_scalability_plot' in plot_paths and os.path.exists(plot_paths['load_scalability_plot']):
            doc.add_heading('Load Scalability Analysis', level=2)
            doc.add_picture(plot_paths['load_scalability_plot'], width=Inches(6))
            
            load_para = doc.add_paragraph("This visualization analyzes how system performance metrics (throughput and response time) ")
            load_para.add_run("change with increasing load levels. It helps identify the saturation point where adding more load ")
            load_para.add_run("no longer improves throughput and may degrade response time.")
            
            doc.add_paragraph()
            key_indicators = doc.add_paragraph()
            key_indicators.add_run("Key Indicators:").bold = True
            
            linear_growth = doc.add_paragraph(style='List Bullet')
            linear_growth.add_run("Linear Growth Region: ").bold = True
            linear_growth.add_run("The initial section where throughput increases linearly with load indicates efficient resource utilization.")
            
            saturation = doc.add_paragraph(style='List Bullet')
            saturation.add_run("Saturation Point: ").bold = True
            saturation.add_run("The load level where throughput begins to level off, indicating system resources becoming fully utilized.")
            
            degradation = doc.add_paragraph(style='List Bullet')
            degradation.add_run("Degradation Region: ").bold = True
            degradation.add_run("The area where throughput may decrease while response time continues to increase, indicating system overload.")
            
            optimal_point = doc.add_paragraph(style='List Bullet')
            optimal_point.add_run("Optimal Load Point: ").bold = True
            optimal_point.add_run("The load level that provides the best balance between high throughput and acceptable response time.")
            
            doc.add_paragraph()
            interpretation = doc.add_paragraph()
            interpretation.add_run("Understanding these patterns helps in capacity planning and identifying the optimal operating ")
            interpretation.add_run("range for your system. It can also guide hardware provisioning decisions and performance tuning efforts.")
        
        # Load Capacity Model if available
        if plot_paths and 'capacity_model_plot' in plot_paths and os.path.exists(plot_paths['capacity_model_plot']):
            doc.add_heading('Load Capacity Model', level=2)
            doc.add_picture(plot_paths['capacity_model_plot'], width=Inches(6))
            
            capacity_para = doc.add_paragraph("This visualization presents a predictive capacity model showing the theoretical maximum ")
            capacity_para.add_run("throughput capacity of the system and the limiting factors that constrain it. The model applies the ")
            capacity_para.add_run("Universal Scalability Law to predict system behavior under various load conditions and resource configurations.")
            
            doc.add_paragraph()
            model_components = doc.add_paragraph()
            model_components.add_run("Model Components:").bold = True
            
            theoretical_max = doc.add_paragraph(style='List Bullet')
            theoretical_max.add_run("Theoretical Maximum: ").bold = True
            theoretical_max.add_run("The asymptotic maximum throughput the system can achieve under ideal conditions.")
            
            contention = doc.add_paragraph(style='List Bullet')
            contention.add_run("Contention Factor (α): ").bold = True
            contention.add_run("Represents resource conflicts and serialization that limit scalability.")
            
            coherency = doc.add_paragraph(style='List Bullet')
            coherency.add_run("Coherency Factor (β): ").bold = True
            coherency.add_run("Represents coordination overhead that causes performance to eventually decrease with added load.")
            
            optimal_op = doc.add_paragraph(style='List Bullet')
            optimal_op.add_run("Optimal Operating Point: ").bold = True
            optimal_op.add_run("The load level that maximizes throughput before coherency costs cause degradation.")
            
            doc.add_paragraph()
            model_value = doc.add_paragraph()
            model_value.add_run("The capacity model helps predict how the system will behave beyond the tested load levels, enabling ")
            model_value.add_run("more accurate capacity planning and infrastructure sizing decisions. By understanding the contention ")
            model_value.add_run("and coherency factors, targeted optimizations can be made to improve overall system scalability.")
            cost_model = lambda x: x
            costs = [cost_model(r) for r in resource_levels]
            throughput_per_cost = [t / c if c > 0 else 0 for t, c in zip(throughputs, costs)]
            most_efficient_idx = throughput_per_cost.index(max(throughput_per_cost)) if throughput_per_cost else 0
            
            cost_obs = doc.add_paragraph(style='List Bullet')
            cost_obs.add_run(f"Most cost-efficient configuration: ").bold = True
            cost_obs.add_run(f"{resource_levels[most_efficient_idx]} resources")
            
            if most_efficient_idx == len(resource_levels) - 1:
                cost_rec = doc.add_paragraph(style='List Bullet')
                cost_rec.add_run("Recommendation: ").bold = True
                cost_rec.add_run("Consider testing with even more resources as cost efficiency is still increasing")
            elif most_efficient_idx == 0:
                cost_rec = doc.add_paragraph(style='List Bullet')
                cost_rec.add_run("Recommendation: ").bold = True
                cost_rec.add_run("Consider testing with fewer resources as cost efficiency is decreasing")
            else:
                cost_rec = doc.add_paragraph(style='List Bullet')
                cost_rec.add_run("Recommendation: ").bold = True
                cost_rec.add_run(f"{resource_levels[most_efficient_idx]} resources provides the best balance of performance and cost")
    
    # Add Theoretical Projections section if available
    if scalability_analysis and scalability_analysis.get('insufficient_data', False):
        theoretical_projections = scalability_analysis.get('theoretical_projections', {})
        if theoretical_projections:
            doc.add_heading('Theoretical Projections Based on Limited Data', 1)
            
            projection_para = doc.add_paragraph(
                "Since only a limited number of resource levels were tested, theoretical projections "
                "have been generated based on the available data points. These projections should be considered "
                "estimates that need validation with additional measurements."
            )
            
            # Amdahl's Law projections
            if 'amdahl' in theoretical_projections:
                p = theoretical_projections['amdahl'].get('parallelizable_fraction', 0)
                max_speedup = theoretical_projections['amdahl'].get('estimated_max_speedup', 0)
                efficiency = theoretical_projections['amdahl'].get('observed_efficiency', 0)
                
                doc.add_heading("Amdahl's Law Projection", 2)
                
                amdahl_para = doc.add_paragraph()
                amdahl_para.add_run(f"Based on the observed speedup efficiency of {efficiency:.2%}, Amdahl's Law suggests that ")
                amdahl_para.add_run(f"{p*100:.1f}% of your workload is parallelizable. ").bold = True
                
                doc.add_heading("Implications:", 3)
                
                implications = doc.add_paragraph(style='List Bullet')
                implications.add_run(f"Theoretical maximum speedup limit: ")
                implications.add_run(f"{max_speedup:.1f}x").bold = True
                
                implications = doc.add_paragraph(style='List Bullet')
                implications.add_run(f"As you add more resources beyond those tested, speedup will approach but never exceed this limit")
            
            # Gustafson's Law projections
            if 'gustafson' in theoretical_projections:
                p = theoretical_projections['gustafson'].get('scalable_fraction', 0)
                
                doc.add_heading("Gustafson's Law Projection", 2)
                
                gustafson_para = doc.add_paragraph()
                gustafson_para.add_run(f"Gustafson's Law suggests that ")
                gustafson_para.add_run(f"{p*100:.1f}% of your workload scales with additional resources. ").bold = True
                
                doc.add_heading("Implications:", 3)
                
                implications = doc.add_paragraph(style='List Bullet')
                implications.add_run(f"If workload size increases with resource count, your system may achieve better scaling than Amdahl's Law predicts")
    
    # Add Optimization Suggestions section if available
    if scalability_analysis and 'optimization_suggestions' in scalability_analysis:
        suggestions = scalability_analysis.get('optimization_suggestions', [])
        if suggestions:
            doc.add_heading('Optimization Suggestions', 1)
            
            for suggestion in suggestions:
                suggestion_para = doc.add_paragraph(style='List Bullet')
                suggestion_para.add_run(suggestion)
    
    # Add Understanding Scalability Models section
    doc.add_heading('Understanding Scalability Models', 1)
    
    model_para = doc.add_paragraph()
    model_para.add_run("This section provides educational information about the different scalability models used in the analysis.")
    
    # Add explanation of each model
    amdahl_explain = doc.add_paragraph(style='List Bullet')
    amdahl_explain.add_run("Amdahl's Law: ").bold = True
    amdahl_explain.add_run("Shows diminishing returns as resources increase due to serial portions of work. "
                         "Higher 'p' values indicate more parallelizable workloads and better scaling potential.")
    
    gustafson_explain = doc.add_paragraph(style='List Bullet')
    gustafson_explain.add_run("Gustafson's Law: ").bold = True
    gustafson_explain.add_run("Shows more optimistic scaling when problem size grows with resources. "
                            "Particularly relevant for workloads that can expand to use available resources.")
    
    usl_explain = doc.add_paragraph(style='List Bullet')
    usl_explain.add_run("Universal Scalability Law: ").bold = True
    usl_explain.add_run("Accounts for both contention (σ) and coherency delays (κ). "
                      "Predicts eventual performance decline at high resource levels due to increasing coordination costs.")
    
    # Add system-specific observation based on observed data
    if len(resource_levels) >= 2:
        doc.add_paragraph()
        system_match = doc.add_paragraph()
        
        # Calculate observed scaling efficiency
        r_ratio = resource_levels[-1] / resource_levels[0]
        speedup = speedups[-1]
        efficiency = speedup / r_ratio if r_ratio > 0 else 0
        
        system_match.add_run("Based on your observed data points, your system currently exhibits: ").bold = True
        if efficiency > 0.95:
            system_match.add_run(f"{efficiency:.1%} scaling efficiency, indicating near-linear scaling. "
                               f"This excellent scaling suggests minimal serial bottlenecks and good resource utilization.")
        elif efficiency > 0.8:
            system_match.add_run(f"{efficiency:.1%} scaling efficiency, indicating good scaling with some overhead. "
                               f"This suggests some serial portions but generally good parallelization.")
        elif efficiency > 0.6:
            system_match.add_run(f"{efficiency:.1%} scaling efficiency, indicating moderate scaling with significant overhead. "
                               f"This suggests substantial serial portions or contention limiting scalability.")
        else:
            system_match.add_run(f"{efficiency:.1%} scaling efficiency, indicating poor scaling. "
                               f"This suggests major bottlenecks that severely limit parallel execution.")
    
    # Save the document
    report_path = os.path.join(output_dir, 'scalability_report.docx')
    doc.save(report_path)
    
    print(f"Enhanced DOCX report saved to {report_path}")
    return report_path


# Main entry point for direct script execution
if __name__ == "__main__":
    import argparse
    from scalability_core import analyze_jtl, create_output_dir
    
    parser = argparse.ArgumentParser(description='Scalability Reporting Module - DOCX Generator')
    parser.add_argument('--files', nargs='+', required=True, help='JTL files to analyze')
    parser.add_argument('--levels', nargs='+', type=int, required=True, 
                        help='Resource levels corresponding to files')
    parser.add_argument('--output-dir', type=str, help='Output directory for report')
    parser.add_argument('--plot-dir', type=str, help='Directory containing plot images to include')
    
    args = parser.parse_args()
    
    if len(args.files) != len(args.levels):
        print("Error: Number of files must match number of resource levels")
        exit(1)
    
    # Create output directory
    output_dir = args.output_dir or create_output_dir()
    
    # Analyze files
    analysis_results = []
    
    for file_path, level in zip(args.files, args.levels):
        metrics = analyze_jtl(file_path)
        if metrics:
            analysis_results.append({
                'file': file_path,
                'resource_level': level,
                'metrics': metrics
            })
    
    # Load plot paths if plot directory provided
    plot_paths = None
    if args.plot_dir:
        plot_paths = {}
        for plot_type in ['throughput', 'response_time', 'speedup', 'models_comparison', 
                         'theoretical_projections', 'model_characteristics']:
            potential_path = os.path.join(args.plot_dir, f"{plot_type}_vs_resource.png")
            if os.path.exists(potential_path):
                plot_paths[plot_type] = potential_path
            
            # Try alternative naming patterns
            if plot_type not in plot_paths:
                alternative_path = os.path.join(args.plot_dir, f"{plot_type}.png")
                if os.path.exists(alternative_path):
                    plot_paths[plot_type] = alternative_path
    
    # Generate report
    if analysis_results:
        generate_docx_report(analysis_results, output_dir, plot_paths)
    else:
        print("Error: No valid data to generate report")
