import pandas as pd
import matplotlib.pyplot as plt
import os
import argparse
from datetime import datetime
import xml.etree.ElementTree as ET
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown2
import numpy as np
from scalability_models import (
    amdahls_law, fit_amdahls_law,
    gustafsons_law, fit_gustafsons_law,
    universal_scalability_law, fit_universal_scalability_law,
    plot_scalability_models, plot_theoretical_projections,
    interpret_scalability_results, suggest_optimizations
)
import json

def analyze_jtl(file_path):
    """Analyzes a single JTL file (XML format) and returns key performance metrics."""
    samples = []
    try:
        # Use iterparse for robust parsing of potentially large or malformed files
        for event, elem in ET.iterparse(file_path, events=('end',)):
            if elem.tag == 'sample':
                try:
                    samples.append({
                        'timeStamp': int(elem.get('ts')),
                        'elapsed': int(elem.get('t')),
                        'success': elem.get('s') == 'true'
                    })
                except (TypeError, ValueError):
                    # This sample is missing attributes, skip it.
                    pass
                finally:
                    # Free up memory
                    elem.clear()
    except ET.ParseError as e:
        print(f"Warning: XML parsing error in {file_path}: {e}. Processing any data found before the error.")

    if not samples:
        print(f"Error: No valid sample data could be read from {file_path}.")
        return None

    df = pd.DataFrame(samples)

    # Convert timestamp to datetime
    df['timeStamp'] = pd.to_datetime(df['timeStamp'], unit='ms')

    # Calculate test duration
    if df.empty:
        return None
    start_time = df['timeStamp'].min()
    end_time = df['timeStamp'].max()
    duration_seconds = (end_time - start_time).total_seconds()

    # Calculate metrics
    total_requests = len(df)
    throughput = total_requests / duration_seconds if duration_seconds > 0 else 0
    avg_response_time = df['elapsed'].mean()
    error_count = df[df['success'] == False].shape[0]
    error_percentage = (error_count / total_requests) * 100 if total_requests > 0 else 0

    return {
        'total_requests': total_requests,
        'duration_seconds': duration_seconds,
        'throughput': throughput,
        'avg_response_time': avg_response_time,
        'error_percentage': error_percentage
    }

def perform_scalability_analysis(results):
    """Perform in-depth scalability analysis using various scalability models.
    
    Args:
        results (list): List of test results with metrics and resource levels
        
    Returns:
        dict: Analysis results including model parameters and interpretations
    """
    # Sort results by resource level for consistent analysis
    sorted_results = sorted(results, key=lambda x: x['resource_level'])
    
    # Extract resource levels and performance metrics
    resource_levels = [r['resource_level'] for r in sorted_results]
    throughputs = [r['metrics']['throughput'] for r in sorted_results]
    
    # Calculate speedups relative to the baseline
    baseline_throughput = throughputs[0]
    speedups = [t / baseline_throughput for t in throughputs]
    
    # Initialize results dictionary
    analysis = {
        'resource_levels': resource_levels,
        'actual_speedups': speedups,
        'models': {},
        'interpretations': {},
        'optimization_suggestions': [],
        'theoretical_projections': {},
        'insufficient_data': False
    }
    
    # Check if we have enough data points (at least 3 for fitting)
    if len(resource_levels) >= 3:
        # Fit Amdahl's Law
        try:
            p_amdahl, error_amdahl = fit_amdahls_law(resource_levels, speedups)
            analysis['models']['amdahl'] = p_amdahl
            analysis['model_errors'] = {'amdahl': error_amdahl}
        except Exception as e:
            print(f"Warning: Could not fit Amdahl's Law: {e}")
        
        # Fit Gustafson's Law
        try:
            p_gustafson, error_gustafson = fit_gustafsons_law(resource_levels, speedups)
            analysis['models']['gustafson'] = p_gustafson
            analysis['model_errors']['gustafson'] = error_gustafson
        except Exception as e:
            print(f"Warning: Could not fit Gustafson's Law: {e}")
        
        # Fit Universal Scalability Law
        try:
            sigma, kappa, error_usl = fit_universal_scalability_law(resource_levels, speedups)
            analysis['models']['usl'] = (sigma, kappa)
            analysis['model_errors']['usl'] = error_usl
        except Exception as e:
            print(f"Warning: Could not fit Universal Scalability Law: {e}")
        
        # Generate interpretations if models were successfully fit
        if analysis['models']:
            analysis['interpretations'] = interpret_scalability_results(analysis['models'])
            analysis['optimization_suggestions'] = suggest_optimizations(analysis['models'], analysis['interpretations'])
    else:
        # Not enough data points for fitting, but we can still provide theoretical projections
        print("Warning: Not enough data points to fit scalability models. At least 3 resource levels are required.")
        analysis['insufficient_data'] = True
        
        # Calculate theoretical projections based on the limited data we have
        # For 2 data points, we can estimate p for Amdahl's and Gustafson's Laws
        if len(resource_levels) == 2:
            # Calculate observed efficiency to estimate Amdahl's parallelizable fraction
            r_ratio = resource_levels[1] / resource_levels[0]
            speedup = speedups[1]
            efficiency = speedup / r_ratio
            
            # Estimate Amdahl's Law parameter (p) by solving for the observed speedup
            # 1/((1-p) + p/n) = observed speedup, solve for p
            try:
                n = r_ratio
                observed_speedup = speedup
                # p = (n - observed_speedup) / (n * observed_speedup - observed_speedup)
                
                # A more robust approach (using scipy.optimize)
                def amdahl_error(p):
                    predicted = amdahls_law(p, n)
                    return abs(predicted - observed_speedup)
                
                # Find p that minimizes error
                from scipy.optimize import minimize_scalar
                result = minimize_scalar(amdahl_error, bounds=(0, 1), method='bounded')
                estimated_p_amdahl = result.x
                
                analysis['theoretical_projections']['amdahl'] = {
                    'parallelizable_fraction': estimated_p_amdahl,
                    'observed_efficiency': efficiency,
                    'estimated_max_speedup': 1/(1-estimated_p_amdahl) if estimated_p_amdahl < 1 else float('inf')
                }
                
                # Also provide a rough estimate for Gustafson's Law parameter
                # Using the same p as Amdahl's for simplicity
                analysis['theoretical_projections']['gustafson'] = {
                    'scalable_fraction': estimated_p_amdahl
                }
                
                # For USL, we need at least 3 points, but we can use typical values
                # for illustration purposes only
                analysis['theoretical_projections']['usl'] = {
                    'typical_sigma': 0.1,  # A moderate contention value
                    'typical_kappa': 0.01   # A typical coherency delay value
                }
                
            except Exception as e:
                print(f"Warning: Could not calculate theoretical projections: {e}")
    
    return analysis

def generate_report(results, output_dir):
    """Generates reports in Markdown, HTML, and DOCX formats, including plots."""
    if not results:
        print("No results to generate a report.")
        return

    # Sort results by resource level for consistent reporting
    sorted_results = sorted(results, key=lambda x: x['resource_level'])
    baseline = sorted_results[0]

    # --- Generate Plots --- (needed for all report types)
    resource_levels = [r['resource_level'] for r in sorted_results]
    throughputs = [r['metrics']['throughput'] for r in sorted_results]
    avg_response_times = [r['metrics']['avg_response_time'] for r in sorted_results]
    speedups = [t / throughputs[0] for t in throughputs]
    
    # Perform scalability analysis
    scalability_analysis = perform_scalability_analysis(results)

    plot_paths = {}
    # Plot 1: Throughput vs. Resource Level (Enhanced with annotations)
    plt.figure(figsize=(10, 6))
    plt.plot(resource_levels, throughputs, marker='o', linewidth=2, markersize=8)
    
    # Add trend line for better visualization
    z = np.polyfit(resource_levels, throughputs, 1)
    p = np.poly1d(z)
    plt.plot(resource_levels, p(resource_levels), "--", color='orange', alpha=0.7, 
             label=f"Trend: {z[0]:.2f}x + {z[1]:.2f}")
    
    # Add annotations for key points
    max_throughput_idx = throughputs.index(max(throughputs))
    plt.annotate(f"Max: {max(throughputs):.2f} req/s", 
                 xy=(resource_levels[max_throughput_idx], max(throughputs)),
                 xytext=(5, 10), textcoords='offset points',
                 arrowprops=dict(arrowstyle="->", connectionstyle="arc3,rad=.2"))
    
    plt.title('Throughput vs. Resource Level', fontsize=14, fontweight='bold')
    plt.xlabel('Resource Level (e.g., number of nodes)', fontsize=12)
    plt.ylabel('Throughput (requests/sec)', fontsize=12)
    plt.grid(True, alpha=0.3)
    plt.legend()
    plot_paths['throughput'] = os.path.join(output_dir, 'throughput_vs_resource.png')
    plt.savefig(plot_paths['throughput'], dpi=150)
    plt.close()

    # Plot 2: Avg Response Time vs. Resource Level (Enhanced)
    plt.figure(figsize=(10, 6))
    plt.plot(resource_levels, avg_response_times, marker='o', color='r', linewidth=2, markersize=8)
    
    # Add horizontal line for SLA if it exists
    ideal_response_time = min(avg_response_times)
    plt.axhline(y=ideal_response_time*1.5, color='r', linestyle='--', alpha=0.5, 
                label=f"SLA Threshold (150% of best: {ideal_response_time*1.5:.2f} ms)")
    
    # Add annotations for min/max points
    min_rt_idx = avg_response_times.index(min(avg_response_times))
    max_rt_idx = avg_response_times.index(max(avg_response_times))
    
    plt.annotate(f"Best: {min(avg_response_times):.2f} ms", 
                 xy=(resource_levels[min_rt_idx], min(avg_response_times)),
                 xytext=(5, 10), textcoords='offset points',
                 arrowprops=dict(arrowstyle="->", connectionstyle="arc3,rad=.2"))
                 
    if max_rt_idx != min_rt_idx:  # Only add if different from min
        plt.annotate(f"Worst: {max(avg_response_times):.2f} ms", 
                    xy=(resource_levels[max_rt_idx], max(avg_response_times)),
                    xytext=(5, 10), textcoords='offset points',
                    arrowprops=dict(arrowstyle="->", connectionstyle="arc3,rad=.2"))
    
    plt.title('Average Response Time vs. Resource Level', fontsize=14, fontweight='bold')
    plt.xlabel('Resource Level (e.g., number of nodes)', fontsize=12)
    plt.ylabel('Average Response Time (ms)', fontsize=12)
    plt.grid(True, alpha=0.3)
    plt.legend()
    plot_paths['response_time'] = os.path.join(output_dir, 'response_time_vs_resource.png')
    plt.savefig(plot_paths['response_time'], dpi=150)
    plt.close()
    
    # Plot 3: Throughput Speedup vs. Resource Level (Enhanced with model comparisons)
    plt.figure(figsize=(10, 6))
    plt.plot(resource_levels, speedups, marker='o', color='g', linewidth=2, markersize=8, label='Actual Speedup')
    plt.plot(resource_levels, [r / resource_levels[0] for r in resource_levels], linestyle='--', 
             color='gray', label='Ideal Linear Speedup')
    
    # Add model lines if available
    smooth_x = np.linspace(min(resource_levels), max(resource_levels) * 1.2, 50)
    if 'amdahl' in scalability_analysis['models']:
        p = scalability_analysis['models']['amdahl']
        amdahl_speedups = amdahls_law(p, smooth_x)
        plt.plot(smooth_x, amdahl_speedups, 'r-', label=f"Amdahl's Law (p={p:.3f})")
    
    if 'gustafson' in scalability_analysis['models']:
        p = scalability_analysis['models']['gustafson']
        gustafson_speedups = gustafsons_law(p, smooth_x)
        plt.plot(smooth_x, gustafson_speedups, 'b-', label=f"Gustafson's Law (p={p:.3f})")
    
    if 'usl' in scalability_analysis['models']:
        sigma, kappa = scalability_analysis['models']['usl']
        usl_speedups = universal_scalability_law(sigma, kappa, smooth_x)
        plt.plot(smooth_x, usl_speedups, 'm-', label=f"USL (σ={sigma:.3f}, κ={kappa:.3f})")
    
    # Add annotations for efficiency
    max_efficiency_idx = 0
    max_efficiency = 0
    for i, (r, s) in enumerate(zip(resource_levels, speedups)):
        efficiency = s / (r / resource_levels[0])
        if efficiency > max_efficiency:
            max_efficiency = efficiency
            max_efficiency_idx = i
    
    plt.annotate(f"Best Efficiency: {max_efficiency:.2%}", 
                xy=(resource_levels[max_efficiency_idx], speedups[max_efficiency_idx]),
                xytext=(5, -25), textcoords='offset points',
                arrowprops=dict(arrowstyle="->", connectionstyle="arc3,rad=.2"))
    
    plt.title('Scalability Analysis: Speedup vs. Resource Level', fontsize=14, fontweight='bold')
    plt.xlabel('Resource Level (e.g., number of nodes)', fontsize=12)
    plt.ylabel('Throughput Speedup', fontsize=12)
    plt.grid(True, alpha=0.3)
    plt.legend(loc='best')
    plot_paths['speedup'] = os.path.join(output_dir, 'speedup_vs_resource.png')
    plt.savefig(plot_paths['speedup'], dpi=150)
    plt.close()
    
    # Generate combined scalability models plot if we have enough data points
    if len(resource_levels) >= 3 and scalability_analysis['models']:
        model_plot_paths = plot_scalability_models(
            resource_levels, 
            scalability_analysis['actual_speedups'],
            output_dir,
            scalability_analysis['models']
        )
        plot_paths.update(model_plot_paths)
    
    # Generate theoretical projections plots even with limited data points
    if scalability_analysis['insufficient_data'] and 'theoretical_projections' in scalability_analysis:
        projection_plot_paths = plot_theoretical_projections(
            resource_levels,
            speedups,
            output_dir,
            scalability_analysis['theoretical_projections'],
            max_projection=16  # Project up to 16x resources
        )
        plot_paths.update(projection_plot_paths)
    
    # Always generate comparative model plots for educational purposes
    comparative_plot_paths = plot_theoretical_projections(
        resource_levels,
        speedups,
        output_dir
    )
    plot_paths.update(comparative_plot_paths)
    
    print(f"Enhanced plots saved to {output_dir}")

    # --- Generate Markdown Content ---
    report_md = f"# Scalability Analysis Report\n\n**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
    
    # Executive Summary section
    report_md += "## Executive Summary\n\n"
    # Calculate overall metrics
    max_throughput = max(throughputs)
    max_throughput_level = resource_levels[throughputs.index(max_throughput)]
    min_response_time = min(avg_response_times)
    min_rt_level = resource_levels[avg_response_times.index(min_response_time)]
    max_speedup = max(speedups)
    max_speedup_level = resource_levels[speedups.index(max_speedup)]
    
    # Add executive summary text
    report_md += f"This report analyzes the scalability characteristics of the system under test across {len(resource_levels)} different resource levels "  
    report_md += f"(from {min(resource_levels)} to {max(resource_levels)}).\n\n"
    report_md += "**Key Findings:**\n\n"
    report_md += f"- **Maximum Throughput:** {max_throughput:.2f} requests/sec achieved at resource level {max_throughput_level}\n"
    report_md += f"- **Best Response Time:** {min_response_time:.2f} ms achieved at resource level {min_rt_level}\n"
    report_md += f"- **Maximum Speedup:** {max_speedup:.2f}x achieved at resource level {max_speedup_level} compared to baseline\n"
    
    # Add scalability model summaries if available
    if 'amdahl' in scalability_analysis['models'] and 'interpretations' in scalability_analysis:
        p = scalability_analysis['models']['amdahl']
        max_theoretical = scalability_analysis['interpretations']['amdahl']['max_theoretical_speedup']
        report_md += f"- **Amdahl's Law Analysis:** {p*100:.1f}% of workload is parallelizable (theoretical max speedup: {max_theoretical:.1f}x)\n"
    if 'gustafson' in scalability_analysis['models'] and 'interpretations' in scalability_analysis:
        p = scalability_analysis['models']['gustafson']
        report_md += f"- **Gustafson's Law Analysis:** {p*100:.1f}% of workload scales with additional resources\n"
    if 'usl' in scalability_analysis['models'] and 'interpretations' in scalability_analysis:
        sigma, kappa = scalability_analysis['models']['usl']
        peak_n = scalability_analysis['interpretations']['usl'].get('peak_concurrency', float('inf'))
        if peak_n < float('inf'):
            report_md += f"- **Universal Scalability Law:** System performance predicted to peak at {peak_n:.1f} resources\n"
    
    # Add optimization suggestions if available
    if 'optimization_suggestions' in scalability_analysis and scalability_analysis['optimization_suggestions']:
        report_md += "\n**Optimization Suggestions:**\n\n"
        for i, suggestion in enumerate(scalability_analysis['optimization_suggestions'], 1):
            report_md += f"{i}. {suggestion}\n"
    
    # Detailed Performance Metrics section
    report_md += "\n## Detailed Performance Metrics\n\n"
    report_md += "| Resource Level | Throughput (req/s) | Avg Response Time (ms) | Error % |\n"
    report_md += "|---|---|---|---|\n"
    for result in sorted_results:
        report_md += f"| {result['resource_level']} | {result['metrics']['throughput']:.2f} | {result['metrics']['avg_response_time']:.2f} | {result['metrics']['error_percentage']:.2f} |\n"
    
    # Basic Scalability Metrics section
    report_md += "\n## Basic Scalability Metrics\n\n"
    report_md += "| Resource Level | Throughput Speedup | Scalability Efficiency |\n"
    report_md += "|---|---|---|\n"
    for result in sorted_results:
        resource_ratio = result['resource_level'] / baseline['resource_level']
        speedup = result['metrics']['throughput'] / baseline['metrics']['throughput']
        efficiency = speedup / resource_ratio if resource_ratio > 0 else 0
        report_md += f"| {result['resource_level']} | {speedup:.2f}x | {efficiency:.2%} |\n"
        
    # Advanced Scalability Analysis section (if models are available)
    if scalability_analysis['models'] and 'interpretations' in scalability_analysis:
        report_md += "\n## Advanced Scalability Analysis\n\n"
        
        # Amdahl's Law section
        if 'amdahl' in scalability_analysis['models']:
            p = scalability_analysis['models']['amdahl']
            assessment = scalability_analysis['interpretations']['amdahl']['assessment']
            report_md += "### Amdahl's Law Analysis\n\n"
            report_md += f"Amdahl's Law models the theoretical speedup of a system when parts of it are parallelizable. "  
            report_md += f"Based on the observed data, **{p*100:.1f}%** of the workload appears to be parallelizable.\n\n"
            report_md += f"**Assessment:** {assessment}\n\n"
            report_md += f"**Maximum Theoretical Speedup:** {scalability_analysis['interpretations']['amdahl']['max_theoretical_speedup']:.1f}x\n\n"
        
        # Gustafson's Law section
        if 'gustafson' in scalability_analysis['models']:
            p = scalability_analysis['models']['gustafson']
            assessment = scalability_analysis['interpretations']['gustafson']['assessment']
            report_md += "### Gustafson's Law Analysis\n\n"
            report_md += f"Gustafson's Law models how systems scale with increased workload sizes. "  
            report_md += f"Based on the observed data, **{p*100:.1f}%** of the workload scales with additional resources.\n\n"
            report_md += f"**Assessment:** {assessment}\n\n"
        
        # Universal Scalability Law section
        if 'usl' in scalability_analysis['models']:
            sigma, kappa = scalability_analysis['models']['usl']
            assessment = scalability_analysis['interpretations']['usl']['assessment']
            peak_n = scalability_analysis['interpretations']['usl']['peak_concurrency']
            
            report_md += "### Universal Scalability Law Analysis\n\n"
            report_md += "The Universal Scalability Law (USL) extends Amdahl's Law by accounting for both contention (σ) and " 
            report_md += "coherency delay (κ) in parallel systems.\n\n"
            report_md += f"**Contention Factor (σ):** {sigma:.3f} - Represents overhead from resource sharing\n"
            report_md += f"**Coherency Delay (κ):** {kappa:.3f} - Represents overhead from data coherency operations\n\n"
            
            if peak_n < float('inf'):
                report_md += f"**Peak Performance Point:** The system performance is predicted to peak at {peak_n:.1f} resources\n\n"
            
            report_md += f"**Assessment:** {assessment}\n\n"
        
        # Optimization Recommendations
        if scalability_analysis['optimization_suggestions']:
            report_md += "### Optimization Recommendations\n\n"
            for suggestion in scalability_analysis['optimization_suggestions']:
                report_md += f"- {suggestion}\n"
            
    # Visual Analysis section
    report_md += "\n## Visual Analysis\n\n"
    report_md += "Refer to the generated plots for visual representation of the scalability characteristics.\n"
    report_md += "Key visualizations include throughput vs. resource level, response time trends, and comparative analysis of " 
    report_md += "actual speedup against theoretical models.\n\n"
    report_md += "The plots provide insights into:\n"
    report_md += "- System scaling behavior across different resource levels\n"
    report_md += "- Deviation from ideal linear scaling\n"
    report_md += "- Efficiency trends and potential bottlenecks\n"
    if 'models_comparison' in plot_paths:
        report_md += "- Comparison between observed data and theoretical scalability models\n"
    
    # Add information about theoretical projections if we have limited data points
    if scalability_analysis['insufficient_data'] and 'theoretical_projections' in scalability_analysis:
        report_md += "\n## Theoretical Projections Based on Limited Data\n\n"
        report_md += "Since only a limited number of resource levels were tested, the following theoretical projections "
        report_md += "have been generated based on the available data points. These projections should be considered "
        report_md += "estimates that need validation with additional measurements.\n\n"
        
        # Add Amdahl's Law projection information
        if 'amdahl' in scalability_analysis['theoretical_projections']:
            p = scalability_analysis['theoretical_projections']['amdahl']['parallelizable_fraction']
            max_speedup = scalability_analysis['theoretical_projections']['amdahl']['estimated_max_speedup']
            efficiency = scalability_analysis['theoretical_projections']['amdahl']['observed_efficiency']
            
            report_md += "### Amdahl's Law Projection\n\n"
            report_md += f"Based on the observed speedup efficiency of {efficiency:.2%}, Amdahl's Law suggests that "
            report_md += f"approximately **{p*100:.1f}%** of your workload is parallelizable.\n\n"
            report_md += f"**Implications:**\n\n"
            report_md += f"- Theoretical maximum speedup limit: **{max_speedup:.1f}x**\n"
            report_md += f"- As you add more resources beyond those tested, speedup will approach but never exceed this limit\n"
            report_md += f"- With this parallelization factor, doubling resources from current levels would yield approximately "
            report_md += f"**{amdahls_law(p, resource_levels[-1]*2):.2f}x** speedup compared to baseline\n\n"
        
        # Add Gustafson's Law projection information
        if 'gustafson' in scalability_analysis['theoretical_projections']:
            p = scalability_analysis['theoretical_projections']['gustafson']['scalable_fraction']
            
            report_md += "### Gustafson's Law Projection\n\n"
            report_md += f"Gustafson's Law suggests that **{p*100:.1f}%** of your workload scales with additional resources.\n\n"
            report_md += f"**Implications:**\n\n"
            report_md += f"- If workload size increases with resource count, your system may achieve better scaling than Amdahl's Law predicts\n"
            report_md += f"- With this scalability factor, doubling resources from current levels could yield approximately "
            report_md += f"**{gustafsons_law(p, resource_levels[-1]*2):.2f}x** speedup compared to baseline (assuming workload grows)\n\n"
        
        # Add USL information
        if 'usl' in scalability_analysis['theoretical_projections']:
            sigma = scalability_analysis['theoretical_projections']['usl']['typical_sigma']
            kappa = scalability_analysis['theoretical_projections']['usl']['typical_kappa']
            
            report_md += "### Universal Scalability Law Considerations\n\n"
            report_md += "The Universal Scalability Law accounts for both contention and coherency delays in parallel systems. "
            report_md += "While more data points are needed for accurate USL fitting, typical systems exhibit contention and coherency factors "
            report_md += "that can help anticipate scaling behavior.\n\n"
            report_md += "**Note:** The theoretical projection plots include USL curves with typical values for illustration purposes.\n\n"
    
    # Add information about comparative model characteristics
    report_md += "\n## Comparative Scalability Models\n\n"
    report_md += "The comparative model characteristics plot illustrates how different theoretical scalability laws behave "
    report_md += "across resource levels. This visualization helps you understand which model might best describe your system's "
    report_md += "scalability characteristics and predict future performance.\n\n"
    report_md += "**Key observations:**\n\n"
    report_md += "- Amdahl's Law shows diminishing returns as resources increase due to serial portions\n"
    report_md += "- Gustafson's Law shows more optimistic scaling when problem size grows with resources\n"
    report_md += "- Universal Scalability Law predicts eventual performance decline due to coherency costs\n"
    report_md += "- The 'p' value in these models represents the proportion of work that can benefit from parallelization\n"

    # --- Save Markdown Report ---
    report_path_md = os.path.join(output_dir, 'scalability_report.md')
    with open(report_path_md, 'w') as f:
        f.write(report_md)
    print(f"Markdown report saved to {report_path_md}")

    # --- Save HTML Report ---
    # We'll create a more enhanced HTML with CSS styling
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <title>Scalability Analysis Report</title>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; max-width: 1200px; margin: 0 auto; padding: 20px; }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #2980b9; margin-top: 30px; }}
        h3 {{ color: #3498db; }}
        table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
        tr:nth-child(even) {{ background-color: #f9f9f9; }}
        .plot-container {{ margin: 30px 0; }}
        .plot-container img {{ max-width: 100%; box-shadow: 0 2px 5px rgba(0,0,0,0.1); }}
        .plot-description {{ font-style: italic; margin: 5px 0 20px 0; color: #555; }}
        .key-finding {{ font-weight: bold; color: #e74c3c; }}
        .assessment {{ background-color: #f8f9fa; padding: 10px; border-left: 5px solid #2980b9; margin: 15px 0; }}
        .optimization {{ background-color: #f8f9fa; padding: 10px; border-left: 5px solid #27ae60; margin: 15px 0; }}
    </style>
</head>
<body>
    <h1>Scalability Analysis Report</h1>
    <p><strong>Date:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    
    <h2>Executive Summary</h2>
    <p>This report analyzes the scalability characteristics of the system under test across {len(resource_levels)} different resource levels 
       (from {min(resource_levels)} to {max(resource_levels)}).</p>
    
    <h3>Key Findings</h3>
    <ul>
"""

    # Add key findings
    max_throughput = max(throughputs)
    max_throughput_level = resource_levels[throughputs.index(max_throughput)]
    min_response_time = min(avg_response_times)
    min_rt_level = resource_levels[avg_response_times.index(min_response_time)]
    max_speedup = max(speedups)
    max_speedup_level = resource_levels[speedups.index(max_speedup)]
    
    html_content += f"""        <li><span class='key-finding'>Maximum Throughput:</span> {max_throughput:.2f} requests/sec achieved at resource level {max_throughput_level}</li>
        <li><span class='key-finding'>Best Response Time:</span> {min_response_time:.2f} ms achieved at resource level {min_rt_level}</li>
        <li><span class='key-finding'>Maximum Speedup:</span> {max_speedup:.2f}x achieved at resource level {max_speedup_level} compared to baseline</li>
"""
    
    # Add model findings if available
    if 'amdahl' in scalability_analysis['models'] and 'interpretations' in scalability_analysis:
        p = scalability_analysis['models']['amdahl']
        max_theoretical = scalability_analysis['interpretations']['amdahl']['max_theoretical_speedup']
        html_content += f"""        <li><span class='key-finding'>Amdahl's Law Analysis:</span> {p*100:.1f}% of workload is parallelizable (theoretical max speedup: {max_theoretical:.1f}x)</li>
"""
    
    if 'gustafson' in scalability_analysis['models'] and 'interpretations' in scalability_analysis:
        p = scalability_analysis['models']['gustafson']
        html_content += f"""        <li><span class='key-finding'>Gustafson's Law Analysis:</span> {p*100:.1f}% of workload scales with additional resources</li>
"""
    
    if 'usl' in scalability_analysis['models'] and 'interpretations' in scalability_analysis:
        sigma, kappa = scalability_analysis['models']['usl']
        peak_n = scalability_analysis['interpretations']['usl'].get('peak_concurrency', float('inf'))
        if peak_n < float('inf'):
            html_content += f"""        <li><span class='key-finding'>Universal Scalability Law:</span> System performance predicted to peak at {peak_n:.1f} resources</li>
"""
    
    html_content += """    </ul>
"""
    
    # Add optimization suggestions if available
    if 'optimization_suggestions' in scalability_analysis and scalability_analysis['optimization_suggestions']:
        html_content += """    <h3>Optimization Suggestions</h3>
    <ol class='optimization'>
"""
        for suggestion in scalability_analysis['optimization_suggestions']:
            html_content += f"""        <li>{suggestion}</li>
"""
        html_content += """    </ol>
"""
    
    # Add detailed metrics table
    html_content += """    <h2>Detailed Performance Metrics</h2>
    <table>
        <tr>
            <th>Resource Level</th>
            <th>Throughput (req/s)</th>
            <th>Avg Response Time (ms)</th>
            <th>Error %</th>
        </tr>
"""
    
    for result in sorted_results:
        html_content += f"""        <tr>
            <td>{result['resource_level']}</td>
            <td>{result['metrics']['throughput']:.2f}</td>
            <td>{result['metrics']['avg_response_time']:.2f}</td>
            <td>{result['metrics']['error_percentage']:.2f}</td>
        </tr>
"""
    
    html_content += """    </table>
    
    <h2>Basic Scalability Metrics</h2>
    <table>
        <tr>
            <th>Resource Level</th>
            <th>Throughput Speedup</th>
            <th>Scalability Efficiency</th>
        </tr>
"""
    
    for result in sorted_results:
        resource_ratio = result['resource_level'] / baseline['resource_level']
        speedup = result['metrics']['throughput'] / baseline['metrics']['throughput']
        efficiency = speedup / resource_ratio if resource_ratio > 0 else 0
        html_content += f"""        <tr>
            <td>{result['resource_level']}</td>
            <td>{speedup:.2f}x</td>
            <td>{efficiency:.2%}</td>
        </tr>
"""
    
    html_content += """    </table>
"""
    
    # Advanced Scalability Analysis section
    if scalability_analysis['models'] and 'interpretations' in scalability_analysis:
        html_content += """    <h2>Advanced Scalability Analysis</h2>
"""
        
        # Amdahl's Law section
        if 'amdahl' in scalability_analysis['models']:
            p = scalability_analysis['models']['amdahl']
            assessment = scalability_analysis['interpretations']['amdahl']['assessment']
            html_content += f"""    <h3>Amdahl's Law Analysis</h3>
    <p>Amdahl's Law models the theoretical speedup of a system when parts of it are parallelizable. 
       Based on the observed data, <strong>{p*100:.1f}%</strong> of the workload appears to be parallelizable.</p>
    <div class='assessment'>
        <p><strong>Assessment:</strong> {assessment}</p>
        <p><strong>Maximum Theoretical Speedup:</strong> {scalability_analysis['interpretations']['amdahl']['max_theoretical_speedup']:.1f}x</p>
    </div>
"""
        
        # Gustafson's Law section
        if 'gustafson' in scalability_analysis['models']:
            p = scalability_analysis['models']['gustafson']
            assessment = scalability_analysis['interpretations']['gustafson']['assessment']
            html_content += f"""    <h3>Gustafson's Law Analysis</h3>
    <p>Gustafson's Law models how systems scale with increased workload sizes. 
       Based on the observed data, <strong>{p*100:.1f}%</strong> of the workload scales with additional resources.</p>
    <div class='assessment'>
        <p><strong>Assessment:</strong> {assessment}</p>
    </div>
"""
        
        # Universal Scalability Law section
        if 'usl' in scalability_analysis['models']:
            sigma, kappa = scalability_analysis['models']['usl']
            assessment = scalability_analysis['interpretations']['usl']['assessment']
            peak_n = scalability_analysis['interpretations']['usl']['peak_concurrency']
            
            html_content += f"""    <h3>Universal Scalability Law Analysis</h3>
    <p>The Universal Scalability Law (USL) extends Amdahl's Law by accounting for both contention (σ) 
       and coherency delay (κ) in parallel systems.</p>
    <p><strong>Contention Factor (σ):</strong> {sigma:.3f} - Represents overhead from resource sharing<br>
       <strong>Coherency Delay (κ):</strong> {kappa:.3f} - Represents overhead from data coherency operations</p>
"""
            
            if peak_n < float('inf'):
                html_content += f"""    <p><strong>Peak Performance Point:</strong> The system performance is predicted to peak at {peak_n:.1f} resources</p>
"""
            
            html_content += f"""    <div class='assessment'>
        <p><strong>Assessment:</strong> {assessment}</p>
    </div>
"""
    
    # Visual Analysis section
    html_content += """    <h2>Visual Analysis</h2>
    <p>The following plots provide visual representation of the scalability characteristics.</p>
    
    <h3>Throughput vs. Resource Level</h3>
    <div class='plot-container'>
        <img src='{0}' alt='Throughput Plot'>
        <p class='plot-description'>This plot shows how the system throughput changes as resources are added. The trend line indicates the overall scaling pattern.</p>
    </div>
    
    <h3>Response Time vs. Resource Level</h3>
    <div class='plot-container'>
        <img src='{1}' alt='Response Time Plot'>
        <p class='plot-description'>This plot illustrates how response times are affected by resource scaling. Lower values indicate better performance, and the SLA threshold line represents a reference point for acceptable response times.</p>
    </div>
    
    <h3>Speedup vs. Resource Level</h3>
    <div class='plot-container'>
        <img src='{2}' alt='Speedup Plot'>
        <p class='plot-description'>This plot compares actual speedup against ideal linear speedup and theoretical models. The gap between actual and ideal lines indicates efficiency loss as resources scale.</p>
    </div>
""".format(
        os.path.basename(plot_paths['throughput']),
        os.path.basename(plot_paths['response_time']),
        os.path.basename(plot_paths['speedup'])
    )
    
    # Add combined models plot if available
    if 'models_comparison' in plot_paths:
        html_content += """    <h3>Scalability Models Comparison</h3>
    <div class='plot-container'>
        <img src='{0}' alt='Models Comparison Plot'>
        <p class='plot-description'>This plot compares the actual speedup with predictions from different scalability laws. The closest model to actual data points indicates which theoretical model best describes the system's scaling behavior.</p>
    </div>
""".format(os.path.basename(plot_paths['models_comparison']))
    
    # Add theoretical projection plots if available
    if 'theoretical_projections' in plot_paths:
        html_content += """    <h3>Theoretical Scalability Projections</h3>
    <div class='plot-container'>
        <img src='{0}' alt='Theoretical Projections Plot'>
        <p class='plot-description'>This plot shows theoretical projections of different scalability models based on observed data. It predicts how the system might scale with additional resources beyond those tested.</p>
    </div>
""".format(os.path.basename(plot_paths['theoretical_projections']))
    
    # Add comparative model characteristics plot if available
    if 'model_characteristics' in plot_paths:
        html_content += """    <h3>Comparative Scalability Model Characteristics</h3>
    <div class='plot-container'>
        <img src='{0}' alt='Model Characteristics Plot'>
        <p class='plot-description'>This educational plot illustrates the fundamental differences between various scalability models with different parameters. It helps identify which theoretical model best describes your system's behavior.</p>
    </div>
""".format(os.path.basename(plot_paths['model_characteristics']))
    
    # Close HTML tags
    html_content += """</body>
</html>"""
    
    # Save HTML report
    report_path_html = os.path.join(output_dir, 'scalability_report.html')
    with open(report_path_html, 'w') as f:
        f.write(html_content)
    print(f"Enhanced HTML report saved to {report_path_html}")

    # --- Save DOCX Report ---
    doc = docx.Document()
    doc.add_heading('Scalability Analysis Report', 0)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Executive Summary Section
    doc.add_heading('Executive Summary', level=1)
    summary_para = doc.add_paragraph()
    summary_para.add_run(f"This report analyzes the scalability characteristics of the system under test across {len(resource_levels)} different resource levels "
                      f"(from {min(resource_levels)} to {max(resource_levels)}).").bold = False
    
    # Key Findings
    doc.add_heading('Key Findings', level=2)
    findings_list = doc.add_paragraph(style='List Bullet')
    max_throughput = max(throughputs)
    max_throughput_level = resource_levels[throughputs.index(max_throughput)]
    findings_list.add_run(f"Maximum Throughput: {max_throughput:.2f} requests/sec achieved at resource level {max_throughput_level}").bold = True
    
    findings_list = doc.add_paragraph(style='List Bullet')
    min_response_time = min(avg_response_times)
    min_rt_level = resource_levels[avg_response_times.index(min_response_time)]
    findings_list.add_run(f"Best Response Time: {min_response_time:.2f} ms achieved at resource level {min_rt_level}").bold = True
    
    findings_list = doc.add_paragraph(style='List Bullet')
    max_speedup = max(speedups)
    max_speedup_level = resource_levels[speedups.index(max_speedup)]
    findings_list.add_run(f"Maximum Speedup: {max_speedup:.2f}x achieved at resource level {max_speedup_level} compared to baseline").bold = True
    
    # Add model findings if available
    if 'amdahl' in scalability_analysis['models'] and 'interpretations' in scalability_analysis:
        findings_list = doc.add_paragraph(style='List Bullet')
        p = scalability_analysis['models']['amdahl']
        max_theoretical = scalability_analysis['interpretations']['amdahl']['max_theoretical_speedup']
        findings_list.add_run(f"Amdahl's Law Analysis: {p*100:.1f}% of workload is parallelizable (theoretical max speedup: {max_theoretical:.1f}x)").bold = True
    
    if 'gustafson' in scalability_analysis['models'] and 'interpretations' in scalability_analysis:
        findings_list = doc.add_paragraph(style='List Bullet')
        p = scalability_analysis['models']['gustafson']
        findings_list.add_run(f"Gustafson's Law Analysis: {p*100:.1f}% of workload scales with additional resources").bold = True
    
    if 'usl' in scalability_analysis['models'] and 'interpretations' in scalability_analysis:
        findings_list = doc.add_paragraph(style='List Bullet')
        sigma, kappa = scalability_analysis['models']['usl']
        peak_n = scalability_analysis['interpretations']['usl'].get('peak_concurrency', float('inf'))
        if peak_n < float('inf'):
            findings_list.add_run(f"Universal Scalability Law: System performance predicted to peak at {peak_n:.1f} resources").bold = True
    
    # Add optimization suggestions
    if 'optimization_suggestions' in scalability_analysis and scalability_analysis['optimization_suggestions']:
        doc.add_heading('Optimization Suggestions', level=2)
        for i, suggestion in enumerate(scalability_analysis['optimization_suggestions'], 1):
            opt_para = doc.add_paragraph(style='List Number')
            opt_para.add_run(suggestion)
    
    # Detailed Metrics Section
    doc.add_heading('Detailed Performance Metrics', level=1)
    # Summary Table
    table_summary = doc.add_table(rows=1, cols=4)
    table_summary.style = 'Table Grid'
    hdr_cells = table_summary.rows[0].cells
    hdr_cells[0].text = 'Resource Level'
    hdr_cells[1].text = 'Throughput (req/s)'
    hdr_cells[2].text = 'Avg Response Time (ms)'
    hdr_cells[3].text = 'Error %'
    for result in sorted_results:
        row_cells = table_summary.add_row().cells
        row_cells[0].text = str(result['resource_level'])
        row_cells[1].text = f"{result['metrics']['throughput']:.2f}"
        row_cells[2].text = f"{result['metrics']['avg_response_time']:.2f}"
        row_cells[3].text = f"{result['metrics']['error_percentage']:.2f}"
    doc.add_paragraph()
    
    # Basic Scalability Metrics Section
    doc.add_heading('Basic Scalability Metrics', level=1)
    table_scalability = doc.add_table(rows=1, cols=3)
    table_scalability.style = 'Table Grid'
    hdr_cells = table_scalability.rows[0].cells
    hdr_cells[0].text = 'Resource Level'
    hdr_cells[1].text = 'Throughput Speedup'
    hdr_cells[2].text = 'Scalability Efficiency'
    for result in sorted_results:
        row_cells = table_scalability.add_row().cells
        resource_ratio = result['resource_level'] / baseline['resource_level']
        speedup = result['metrics']['throughput'] / baseline['metrics']['throughput']
        efficiency = speedup / resource_ratio if resource_ratio > 0 else 0
        row_cells[0].text = str(result['resource_level'])
        row_cells[1].text = f"{speedup:.2f}x"
        row_cells[2].text = f"{efficiency:.2%}"
    doc.add_paragraph()
    
    # Advanced Scalability Analysis section (if models are available)
    if scalability_analysis['models'] and 'interpretations' in scalability_analysis:
        doc.add_heading('Advanced Scalability Analysis', level=1)
        
        # Amdahl's Law section
        if 'amdahl' in scalability_analysis['models']:
            doc.add_heading("Amdahl's Law Analysis", level=2)
            p = scalability_analysis['models']['amdahl']
            assessment = scalability_analysis['interpretations']['amdahl']['assessment']
            
            amdahl_para = doc.add_paragraph()
            amdahl_para.add_run(f"Amdahl's Law models the theoretical speedup of a system when parts of it are parallelizable. "
                               f"Based on the observed data, ")
            amdahl_para.add_run(f"{p*100:.1f}% of the workload appears to be parallelizable.").bold = True
            
            doc.add_paragraph(f"Assessment: {assessment}")
            doc.add_paragraph(f"Maximum Theoretical Speedup: {scalability_analysis['interpretations']['amdahl']['max_theoretical_speedup']:.1f}x")
        
        # Gustafson's Law section
        if 'gustafson' in scalability_analysis['models']:
            doc.add_heading("Gustafson's Law Analysis", level=2)
            p = scalability_analysis['models']['gustafson']
            assessment = scalability_analysis['interpretations']['gustafson']['assessment']
            
            gustafson_para = doc.add_paragraph()
            gustafson_para.add_run(f"Gustafson's Law models how systems scale with increased workload sizes. "
                                 f"Based on the observed data, ")
            gustafson_para.add_run(f"{p*100:.1f}% of the workload scales with additional resources.").bold = True
            
            doc.add_paragraph(f"Assessment: {assessment}")
        
        # Universal Scalability Law section
        if 'usl' in scalability_analysis['models']:
            doc.add_heading("Universal Scalability Law Analysis", level=2)
            sigma, kappa = scalability_analysis['models']['usl']
            assessment = scalability_analysis['interpretations']['usl']['assessment']
            peak_n = scalability_analysis['interpretations']['usl']['peak_concurrency']
            
            usl_para = doc.add_paragraph()
            usl_para.add_run("The Universal Scalability Law (USL) extends Amdahl's Law by accounting for both contention (σ) "
                            "and coherency delay (κ) in parallel systems.\n")
            
            doc.add_paragraph(f"Contention Factor (σ): {sigma:.3f} - Represents overhead from resource sharing")
            doc.add_paragraph(f"Coherency Delay (κ): {kappa:.3f} - Represents overhead from data coherency operations")
            
            if peak_n < float('inf'):
                peak_para = doc.add_paragraph()
                peak_para.add_run(f"Peak Performance Point: The system performance is predicted to peak at {peak_n:.1f} resources").bold = True
            
            doc.add_paragraph(f"Assessment: {assessment}")
    
    # Visual Analysis section
    doc.add_heading('Visual Analysis', level=1)
    doc.add_paragraph("The following plots provide visual representation of the scalability characteristics:")
    
    # Add plot descriptions
    doc.add_heading("Throughput vs. Resource Level", level=2)
    doc.add_picture(plot_paths['throughput'], width=Inches(6.0))
    doc.add_paragraph("This plot shows how the system throughput changes as resources are added. "  
                      "The trend line indicates the overall scaling pattern.")
    
    doc.add_heading("Response Time vs. Resource Level", level=2)
    doc.add_picture(plot_paths['response_time'], width=Inches(6.0))
    doc.add_paragraph("This plot illustrates how response times are affected by resource scaling. "  
                      "Lower values indicate better performance, and the SLA threshold line represents a reference point for acceptable response times.")
    
    doc.add_heading("Speedup vs. Resource Level", level=2)
    doc.add_picture(plot_paths['speedup'], width=Inches(6.0))
    doc.add_paragraph("This plot compares actual speedup against ideal linear speedup and theoretical models. "  
                      "The gap between actual and ideal lines indicates efficiency loss as resources scale.")
    
    # Add combined models plot if available
    if 'models_comparison' in plot_paths:
        doc.add_heading("Scalability Models Comparison", level=2)
        doc.add_picture(plot_paths['models_comparison'], width=Inches(6.0))
        doc.add_paragraph("This plot compares the actual speedup with predictions from different scalability laws. "  
                          "The closest model to actual data points indicates which theoretical model best describes the system's scaling behavior.")
    
    # Add theoretical projection plots
    if 'theoretical_projections' in plot_paths:
        doc.add_heading("Theoretical Scalability Projections", level=2)
        doc.add_picture(plot_paths['theoretical_projections'], width=Inches(6.0))
        projection_para = doc.add_paragraph()
        projection_para.add_run("This plot shows theoretical projections of different scalability models based on observed data. "
                              "It predicts how the system might scale with additional resources beyond those tested.")
        
        # Add detailed analysis of theoretical projections
        if scalability_analysis['insufficient_data'] and 'theoretical_projections' in scalability_analysis:
            doc.add_paragraph()
            doc.add_heading("Projection Analysis", level=3)
            
            # Amdahl's Law projections
            if 'amdahl' in scalability_analysis['theoretical_projections']:
                p = scalability_analysis['theoretical_projections']['amdahl']['parallelizable_fraction']
                max_speedup = scalability_analysis['theoretical_projections']['amdahl']['estimated_max_speedup']
                efficiency = scalability_analysis['theoretical_projections']['amdahl']['observed_efficiency']
                
                amdahl_para = doc.add_paragraph()
                amdahl_para.add_run(f"Amdahl's Law Analysis: ").bold = True
                amdahl_para.add_run(f"Based on the observed speedup efficiency of {efficiency:.2%}, this model suggests that ")
                amdahl_para.add_run(f"{p*100:.1f}% of your workload is parallelizable. ").bold = True
                amdahl_para.add_run(f"This leads to a theoretical maximum speedup limit of {max_speedup:.1f}x "
                                   f"regardless of how many additional resources are added.")
                
                # Add projection for specific resource levels
                resource_projections = doc.add_paragraph(style='List Bullet')
                resource_projections.add_run(f"Doubling current resources would yield approximately ")
                resource_projections.add_run(f"{amdahls_law(p, resource_levels[-1]*2):.2f}x").bold = True
                resource_projections.add_run(f" speedup compared to baseline")
                
                resource_projections = doc.add_paragraph(style='List Bullet')
                resource_projections.add_run(f"10x current resources would yield approximately ")
                resource_projections.add_run(f"{amdahls_law(p, resource_levels[-1]*10):.2f}x").bold = True
                resource_projections.add_run(f" speedup compared to baseline")
            
            # Gustafson's Law projections
            if 'gustafson' in scalability_analysis['theoretical_projections']:
                p = scalability_analysis['theoretical_projections']['gustafson']['scalable_fraction']
                
                gustafson_para = doc.add_paragraph()
                gustafson_para.add_run(f"Gustafson's Law Analysis: ").bold = True
                gustafson_para.add_run(f"This model suggests that ")
                gustafson_para.add_run(f"{p*100:.1f}% of your workload scales with additional resources. ").bold = True
                gustafson_para.add_run(f"This is relevant when your workload size increases proportionally with resources.")
                
                # Add projection for specific resource levels
                resource_projections = doc.add_paragraph(style='List Bullet')
                resource_projections.add_run(f"Doubling current resources could yield approximately ")
                resource_projections.add_run(f"{gustafsons_law(p, resource_levels[-1]*2):.2f}x").bold = True
                resource_projections.add_run(f" speedup if workload grows proportionally")
                
                resource_projections = doc.add_paragraph(style='List Bullet')
                resource_projections.add_run(f"10x current resources could yield approximately ")
                resource_projections.add_run(f"{gustafsons_law(p, resource_levels[-1]*10):.2f}x").bold = True
                resource_projections.add_run(f" speedup if workload grows proportionally")
    
    # Add comparative model characteristics plot
    if 'model_characteristics' in plot_paths:
        doc.add_heading("Comparative Scalability Model Characteristics", level=2)
        doc.add_picture(plot_paths['model_characteristics'], width=Inches(6.0))
        model_char_para = doc.add_paragraph()
        model_char_para.add_run("This educational plot illustrates the fundamental differences between various scalability models with different parameters. "
                              "It helps identify which theoretical model best describes your system's scaling behavior.")
        
        # Add model comparison explanations
        doc.add_heading("Understanding Scalability Models", level=3)
        
        amdahl_explain = doc.add_paragraph(style='List Bullet')
        amdahl_explain.add_run("Amdahl's Law: ").bold = True
        amdahl_explain.add_run("Shows diminishing returns as resources increase due to serial portions of work. "
                              "Higher 'p' values indicate more parallelizable workloads and better scaling potential.")
        
        gustafson_explain = doc.add_paragraph(style='List Bullet')
        gustafson_explain.add_run("Gustafson's Law: ").bold = True
        gustafson_explain.add_run("Shows more optimistic scaling when problem size grows with resources. "
                                 "Particularly relevant for workloads that can expand to use available resources.")
        
        usl_explain = doc.add_paragraph(style='List Bullet')
        usl_explain.add_run("Universal Scalability Law: ").bold = True
        usl_explain.add_run("Accounts for both contention (σ) and coherency delays (κ). "
                          "Predicts eventual performance decline at high resource levels due to increasing coordination costs.")
        
        # Add system-specific observation
        if len(resource_levels) >= 2:
            doc.add_paragraph()
            system_match = doc.add_paragraph()
            
            # Calculate observed scaling efficiency
            r_ratio = resource_levels[-1] / resource_levels[0]
            speedup = speedups[-1]
            efficiency = speedup / r_ratio
            
            system_match.add_run("Based on your observed data points, your system currently exhibits: ").bold = True
            if efficiency > 0.95:
                system_match.add_run(f"{efficiency:.1%} scaling efficiency, indicating near-linear scaling. "
                                    f"This excellent scaling suggests minimal serial bottlenecks and good resource utilization.")
            elif efficiency > 0.8:
                system_match.add_run(f"{efficiency:.1%} scaling efficiency, indicating good scaling with some overhead. "
                                    f"This suggests some serial portions but generally good parallelization.")
            elif efficiency > 0.6:
                system_match.add_run(f"{efficiency:.1%} scaling efficiency, indicating moderate scaling with significant overhead. "
                                    f"This suggests substantial serial portions or contention limiting scalability.")
            else:
                system_match.add_run(f"{efficiency:.1%} scaling efficiency, indicating poor scaling. "
                                    f"This suggests major bottlenecks that severely limit parallel execution.")
    
    
    report_path_docx = os.path.join(output_dir, 'scalability_report.docx')
    doc.save(report_path_docx)
    print(f"Enhanced DOCX report saved to {report_path_docx}")

def main():
    parser = argparse.ArgumentParser(description='Analyze JMeter JTL files for scalability.')
    parser.add_argument('--files', nargs='+', required=True, help='List of JTL file paths.')
    parser.add_argument('--levels', nargs='+', type=int, required=True, help='Corresponding resource levels for each file (e.g., number of nodes).')
    parser.add_argument('--output-dir', type=str, default='sample_analysis_results', help='Directory to save the report and plots.')
    
    args = parser.parse_args()

    if len(args.files) != len(args.levels):
        print("Error: The number of files must match the number of resource levels.")
        return

    # Create a unique output subdirectory for this run
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    run_output_dir = os.path.join(args.output_dir, f'ScalabilityRun_{timestamp}')
    os.makedirs(run_output_dir, exist_ok=True)

    analysis_results = []
    for file_path, level in zip(args.files, args.levels):
        metrics = analyze_jtl(file_path)
        if metrics:
            analysis_results.append({
                'file': file_path,
                'resource_level': level,
                'metrics': metrics
            })

    generate_report(analysis_results, run_output_dir)

if __name__ == '__main__':
    main()
